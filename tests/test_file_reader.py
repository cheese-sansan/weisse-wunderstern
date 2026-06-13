"""file_reader 单元测试。"""
import unittest
import os
import tempfile
import importlib.util
from utils.file_reader import read_file, _contains_latex, _table_to_markdown


class TestFileReader(unittest.TestCase):

    def test_read_txt_file(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("Hello World")
            tmp = f.name
        try:
            result = read_file(tmp)
            self.assertTrue(result["success"])
            self.assertEqual(result["content"], "Hello World")
            self.assertEqual(result["file_type"], ".txt")
        finally:
            os.unlink(tmp)

    def test_read_json_file(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False, encoding="utf-8") as f:
            f.write('{"key": "value"}')
            tmp = f.name
        try:
            result = read_file(tmp)
            self.assertTrue(result["success"])
            self.assertIn("value", result["content"])
        finally:
            os.unlink(tmp)

    def test_file_not_found(self):
        result = read_file("nonexistent_file_xyz.txt")
        self.assertFalse(result["success"])
        self.assertIsNotNone(result["error"])

    def test_contains_latex(self):
        self.assertTrue(_contains_latex("$E=mc^2$ is Einstein's equation"))
        self.assertTrue(_contains_latex("$$\nx^2\n$$"))
        self.assertFalse(_contains_latex("No formulas here"))

    def test_table_to_markdown(self):
        rows = [
            ["Name", "Score"],
            ["Alice", "95"],
            ["Bob", "87"],
        ]
        md = _table_to_markdown(rows)
        self.assertIn("| Name | Score |", md)
        self.assertIn("| --- | --- |", md)
        self.assertIn("| Alice | 95 |", md)

    def test_table_to_markdown_empty(self):
        self.assertEqual(_table_to_markdown([]), "")
        self.assertEqual(_table_to_markdown([["a"]]), "")

    def test_warnings_in_result(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("test")
            tmp = f.name
        try:
            result = read_file(tmp)
            self.assertIn("warnings", result)
            self.assertIsInstance(result["warnings"], list)
            self.assertIn("contains_latex", result)
            self.assertIn("markdown_text", result)
        finally:
            os.unlink(tmp)

    @unittest.skipIf(importlib.util.find_spec("docx") is None, "python-docx not installed")
    def test_read_docx_table_as_markdown_when_dependency_available(self):
        from docx import Document

        fd, tmp = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc = Document()
            doc.add_paragraph("Benchmark Results")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Model"
            table.cell(0, 1).text = "Accuracy"
            table.cell(1, 0).text = "GPT"
            table.cell(1, 1).text = "86.4"
            doc.save(tmp)

            result = read_file(tmp)
            self.assertTrue(result["success"])
            self.assertIn("Benchmark Results", result["markdown_text"])
            self.assertIn("| Model | Accuracy |", result["markdown_text"])
            self.assertIn("| GPT | 86.4 |", result["markdown_text"])
        finally:
            os.unlink(tmp)


if __name__ == "__main__":
    unittest.main()
